"""Gera docs/ENTREGA-CLIENTE.docx a partir do conteúdo de entrega ao cliente."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ENTREGA-CLIENTE.docx"

# Cores institucionais discretas
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x1E, 0x3A, 0x5F)
HEADER_BG = "1E3A5F"
ALT_ROW = "F4F6F8"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                if not p.runs and p.text:
                    run = p.add_run(p.text)
                    p.clear()
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            if header and row_idx == 0:
                set_cell_shading(cell, HEADER_BG)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif not header and row_idx % 2 == 0:
                set_cell_shading(cell, ALT_ROW)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = ACCENT if level == 1 else INK


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = INK


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    style_table(table)
    doc.add_paragraph()


def add_checkbox_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"☐  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_quote_box(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = MUTED


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = ACCENT

    return doc


def build() -> None:
    doc = setup_document()

    # Capa / título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("Documento de entrega")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.name = "Calibri"
    t_run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("Site institucional\nJulio Capetini e Advogado associado")
    s_run.font.size = Pt(14)
    s_run.font.name = "Calibri"
    s_run.font.color.rgb = INK

    doc.add_paragraph()
    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta_data = [
        ("Versão", "1.0"),
        ("Data", "___/___/2026"),
        ("Elaborado por", "_________________________"),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
    style_table(meta, header=False)
    doc.add_paragraph()

    add_para(
        doc,
        "Este documento resume o que foi entregue, como o site funciona no dia a dia, "
        "onde ficam as contas de serviço e quais são as limitações dos planos gratuitos utilizados. "
        "Serve tanto para referência do escritório quanto para um profissional de TI que venha a fazer manutenção no futuro.",
    )

    # --- 1 ---
    add_heading(doc, "1. Resumo em linguagem simples", 1)
    add_para(
        doc,
        "Foi desenvolvido um site institucional (página profissional na internet) para apresentar o escritório, "
        "a equipa, as áreas de atuação e formas de contacto. O visitante pode:",
    )
    for item in [
        "Conhecer o escritório na página inicial;",
        "Enviar mensagem pelo formulário de contacto (a mensagem chega por e-mail à caixa configurada pelo escritório);",
        "Ler artigos publicados em /artigos (conteúdo jurídico ou informativo);",
        "Contactar por telefone, WhatsApp e e-mail, conforme os dados exibidos no site.",
    ]:
        add_bullet(doc, item)

    add_para(
        doc,
        "A gestão de artigos e de utilizadores administradores é feita numa área reservada, acessível apenas com login "
        "e palavra-passe, em endereços que começam por /admin (por exemplo: /admin/login).",
    )
    add_para(
        doc,
        "O site não substitui software de gestão de processos, agenda jurídica, peticionamento eletrónico ou arquivo de clientes. "
        "É uma vitrine e canal de contacto e conteúdo.",
    )

    # --- 2 ---
    add_heading(doc, "2. O que o site inclui (funcionalidades)", 1)
    add_table(
        doc,
        ["Área", "Endereço (exemplo)", "O que faz"],
        [
            ["Página inicial", "/", "Apresentação, equipa, especialidades, atendimento e formulário de contacto"],
            ["Artigos (público)", "/artigos e /artigos/nome-do-artigo", "Listagem e leitura de textos publicados"],
            ["Login administrativo", "/admin/login", "Entrada da equipa autorizada"],
            ["Gestão de artigos", "/admin/artigos", "Criar, editar e publicar artigos"],
            ["Administradores", "/admin/admins", "Criar ou remover contas de quem pode aceder ao /admin"],
            ["Recuperar senha", "/admin/esqueci-senha", "Envio de link por e-mail para redefinir palavra-passe"],
        ],
    )
    add_para(doc, "E-mails automáticos (via Resend):", bold=True)
    add_bullet(doc, "Mensagens enviadas pelo formulário de contacto da página inicial;")
    add_bullet(doc, "E-mails de recuperação de palavra-passe dos administradores.")

    add_para(doc, "O que não está incluído (limitações de produto):", bold=True)
    for item in [
        "Alterar textos da página inicial (nomes, telefones, biografias) pelo painel admin — exige desenvolvedor;",
        "Área de clientes, upload de documentos de processos ou chat interno;",
        "Loja, pagamentos ou agendamento automático de consultas integrado;",
        "Newsletter ou campanhas de marketing em massa (Resend gratuito = e-mails transacionais).",
    ]:
        add_bullet(doc, item)

    # --- 3 ---
    add_heading(doc, "3. Como o site “vive” na internet", 1)
    add_para(doc, "Fluxo simplificado:", bold=True)
    flow = (
        "Visitante → Domínio (.br no Registro.br) → Hospedagem Vercel (gratuita)\n"
        "    ↳ Base de dados Neon (artigos e admins)\n"
        "    ↳ Resend (e-mails do formulário e recuperação de senha)\n"
        "GitHub → deploy automático → Vercel"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)

    for item in [
        "O domínio foi registado no Registro.br e aponta para os servidores da Vercel.",
        "A Vercel executa o site e liga-se à base de dados Neon (utilizadores admin e artigos).",
        "O Resend envia os e-mails do formulário e da recuperação de senha.",
        "O código está no GitHub; cada atualização aprovada pode gerar um novo deploy na Vercel.",
    ]:
        add_numbered(doc, item)

    add_para(
        doc,
        "E-mail central das contas: foi criada uma caixa no Proton Mail para concentrar logins e recuperação de acesso "
        "(recomendado manter ativa e com autenticação em dois fatores).",
    )

    # --- 4 ---
    add_heading(doc, "4. Serviços contratados (planos gratuitos)", 1)
    add_table(
        doc,
        ["Serviço", "Função", "Plano", "Observação"],
        [
            ["Registro.br", "Domínio .br do site", "Pago anualmente", "Renovação obrigatória"],
            ["Vercel", "Hospedagem e publicação", "Hobby (gratuito)", "Deploy ligado ao Git"],
            ["GitHub", "Código-fonte", "Gratuito", "Histórico e colaboração técnica"],
            ["Neon", "Base de dados PostgreSQL", "Free", "Admins + artigos"],
            ["Resend", "E-mails do site", "Free", "Contacto + reset de senha"],
            ["Proton Mail", "E-mail mestre das contas", "Conforme plano", "Organiza acessos"],
        ],
    )
    add_quote_box(
        doc,
        "Importante: Os limites abaixo são dos planos gratuitos e podem ser alterados pelos fornecedores. "
        "Convém rever o painel de cada serviço uma vez por ano ou antes de campanhas com muito tráfego ou e-mails.",
    )

    # --- 5 ---
    add_heading(doc, "5. Limitações dos planos gratuitos", 1)

    add_heading(doc, "5.1 Vercel (hospedagem)", 2)
    add_table(
        doc,
        ["Limite (Hobby)", "Impacto prático"],
        [
            ["~100 GB/mês de transferência", "Muitas visitas ou imagens pesadas podem pausar o projeto até ao mês seguinte"],
            ["1 milhão de execuções de funções/mês", "Normal para escritório; picos anómalos consomem quota"],
            ["100 deploys/dia", "Só relevante com muitas publicações no mesmo dia"],
            ["Uso comercial", "Validar termos Vercel; empresas costumam migrar para Pro (~US$ 20/mês)"],
        ],
    )
    add_para(doc, "Sintomas: site lento, erro ao publicar, quota excedida no painel Vercel.")

    add_heading(doc, "5.2 Neon (base de dados)", 2)
    add_table(
        doc,
        ["Recurso", "Limite típico (Free)", "Impacto prático"],
        [
            ["Armazenamento", "~0,5 GB (500 MB) por projeto", "Muitos artigos longos podem encher o espaço"],
            ["Compute", "100 CU-hours por projeto/mês", "BD “dorme” após ~5 min; 1.º acesso pode demorar segundos"],
            ["Escala máxima", "Até 2 CU (~8 GB RAM)", "Suficiente para este site"],
            ["Egress", "5 GB/mês", "Raro ser problema em site pequeno"],
        ],
    )
    add_para(doc, "Sintomas: erro ao guardar artigo ou login; compute suspenso no painel Neon.")

    add_heading(doc, "5.3 Resend (e-mail)", 2)
    add_table(
        doc,
        ["Limite (Free)", "Impacto prático"],
        [
            ["3 000 e-mails/mês", "Formulário + recuperações de senha"],
            ["100 e-mails/dia", "Dia intenso pode bloquear envios até 24 h"],
            ["1 domínio verificado", "Remetente deve usar domínio configurado"],
            ["Contagem", "Cada destinatário em Para/CC conta separado"],
        ],
    )
    add_para(doc, "Sintomas: formulário com erro; reset de senha não chega; quota esgotada.")

    add_heading(doc, "5.4 Registro.br (domínio)", 2)
    add_table(
        doc,
        ["Aspecto", "Impacto prático"],
        [
            ["Renovação anual", "Domínio expirado → site inacessível no .br"],
            ["DNS", "Alterações no Registro.br ou onde o DNS estiver delegado"],
            ["WHOIS / titular", "Manter contactos atualizados"],
        ],
    )

    add_heading(doc, "5.5 GitHub", 2)
    add_table(
        doc,
        ["Aspecto", "Impacto prático"],
        [
            ["Repositório", "Definir quem tem acesso de leitura/escrita"],
            ["Sem GitHub", "Novo técnico precisa de cópia do código e acesso Vercel/variáveis"],
        ],
    )

    # --- 6 ---
    add_heading(doc, "6. Contas e acessos (preencher em local seguro)", 1)
    add_quote_box(
        doc,
        "Segurança: Não enviar por e-mail sem encriptação. Preferir gestor de palavras-passe (1Password, Bitwarden). "
        "Ativar autenticação em dois fatores (2FA) em todas as contas.",
    )

    sections_cred = [
        ("6.1 E-mail central (Proton Mail)", [("Endereço", ""), ("Utilização", "Login/recuperação Vercel, GitHub, Neon, Resend, Registro.br"), ("2FA ativo?", "☐ Sim   ☐ Não")]),
        ("6.2 Registro.br", [("Domínio", ""), ("Titular / conta", ""), ("Login", ""), ("Data de renovação", "___/___/____"), ("DNS", "Vercel (painel Vercel → Domains)")]),
        ("6.3 Vercel", [("URL do projeto", "https://"), ("Conta / equipa", ""), ("Repositório ligado", ""), ("Domínio de produção", "https://")]),
        ("6.4 GitHub", [("Organização / utilizador", ""), ("Repositório", ""), ("URL", "https://github.com/"), ("Colaboradores", "")]),
        ("6.5 Neon", [("Projeto", ""), ("Região", ""), ("Connection string", "Só gestor de senhas / Vercel POSTGRES_URL"), ("Painel", "https://console.neon.tech")]),
        ("6.6 Resend", [("Domínio verificado", ""), ("Remetente RESEND_FROM_EMAIL", ""), ("CONTACT_TO_EMAIL", ""), ("API Key", "Só Vercel / gestor de senhas")]),
        ("6.7 Área administrativa", [("URL de login", "https://___/admin/login"), ("E-mail(s) admin", ""), ("Novos admins", "Utilizadores em /admin/admins")]),
    ]

    for title_sec, fields in sections_cred:
        add_heading(doc, title_sec, 2)
        t = doc.add_table(rows=len(fields), cols=2)
        t.style = "Table Grid"
        for i, (k, v) in enumerate(fields):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[1].text = v if v else "___________________________"
        style_table(t, header=False)
        doc.add_paragraph()

    add_heading(doc, "Variáveis de ambiente (Vercel — não estão no GitHub)", 2)
    add_table(
        doc,
        ["Variável", "Para quê"],
        [
            ["POSTGRES_URL", "Ligação à base Neon"],
            ["AUTH_SECRET", "Segurança das sessões admin"],
            ["AUTH_URL", "URL pública do site (https://)"],
            ["NEXT_PUBLIC_SITE_URL", "URL canónica (SEO, partilhas)"],
            ["RESEND_API_KEY", "Chave API Resend"],
            ["RESEND_FROM_EMAIL", "Remetente verificado"],
            ["CONTACT_TO_EMAIL", "Caixa que recebe o formulário"],
        ],
    )

    # --- 7 ---
    add_heading(doc, "7. Operação do dia a dia (para o escritório)", 1)
    add_heading(doc, "Publicar ou editar um artigo", 2)
    for step in [
        "Aceder a https://[seu-dominio]/admin/login",
        "Iniciar sessão",
        "Ir a Artigos → criar ou editar",
        "Publicar — o artigo fica visível em /artigos",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "Receber contactos do site", 2)
    add_para(doc, "As mensagens chegam ao e-mail CONTACT_TO_EMAIL. Verificar pasta de spam nos primeiros dias.")

    add_heading(doc, "Alterar telefone, textos ou logotipo", 2)
    add_para(doc, "Requer desenvolvedor: alteração no código e novo deploy via GitHub → Vercel.")

    add_heading(doc, "Adicionar outro administrador", 2)
    for step in [
        "Login em /admin",
        "/admin/admins → criar utilizador",
        "Entregar credenciais de forma segura",
    ]:
        add_numbered(doc, step)

    # --- 8 ---
    add_heading(doc, "8. Manutenção técnica", 1)
    add_table(
        doc,
        ["Tarefa", "Quando", "Como (resumo)"],
        [
            ["Atualizar dependências / bugs", "Conforme necessidade", "git pull, alterações, push → Vercel"],
            ["Alterar estrutura da BD", "Após mudanças no schema", "npm run db:push com POSTGRES_URL correto"],
            ["Primeiro admin (ambiente novo)", "Uma vez", "npm run db:seed (ADMIN_SEED_* só local)"],
            ["Rever variáveis", "Após mudar domínio ou e-mail", "Painel Vercel"],
            ["Renovar domínio", "Antes da data", "Painel Registro.br"],
        ],
    )

    # --- 9 ---
    add_heading(doc, "9. Riscos e evolução futura", 1)
    add_table(
        doc,
        ["Situação", "Recomendação"],
        [
            ["Muito tráfego ou imagens pesadas", "Monitorizar Vercel; otimizar imagens; plano Pro"],
            ["Muitos artigos / textos enormes", "Acompanhar Neon; plano pago perto de 500 MB"],
            [">100 e-mails/dia no Resend", "Upgrade Resend Pro"],
            ["Editar home sem programador", "Orçar CMS — fora do escopo atual"],
            ["Titular das contas sai", "Transferir Vercel, GitHub, Neon, Resend, Registro.br"],
            ["Perda de acesso", "Proton + 2FA + gestor de senhas + códigos de recuperação"],
        ],
    )
    add_para(doc, "Renovações a calendarizar:", bold=True)
    add_checkbox_item(doc, "Domínio Registro.br — vence em: ___/___/____")
    add_checkbox_item(doc, "Rever planos gratuitos (Vercel / Neon / Resend) — anualmente")
    add_checkbox_item(doc, "Confirmar que CONTACT_TO_EMAIL ainda é monitorizado")

    # --- 10 ---
    add_heading(doc, "10. Checklist de entrega", 1)
    for item in [
        "Site abre no domínio com HTTPS",
        "Formulário de contacto envia e-mail de teste",
        "/artigos lista artigos publicados",
        "Login /admin/login funciona",
        "Criar/editar artigo em /admin/artigos",
        "Recuperação de senha (/admin/esqueci-senha) testada",
        "Anexo de credenciais entregue em canal seguro",
        "2FA ativado nas contas críticas",
        "Responsável interno definido para renovar domínio",
    ]:
        add_checkbox_item(doc, item)

    # --- 11 ---
    add_heading(doc, "11. Contacto do desenvolvedor / suporte", 1)
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Nome / empresa", "_________________________"],
            ["E-mail", "_________________________"],
            ["Telefone", "_________________________"],
            ["Escopo de suporte acordado", "_________________________"],
        ],
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "Documento gerado para entrega ao cliente. "
        "Os limites dos fornecedores (Vercel, Neon, Resend) devem ser confirmados nos respetivos painéis na data da leitura."
    )
    fr.font.size = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = MUTED
    fr.font.name = "Calibri"

    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    build()
